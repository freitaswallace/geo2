#!/usr/bin/env python3
"""
Script para automatizar o processamento de Memoriais Descritivos

Modos de opera√ß√£o:
1. Modo Normal: Processa PDF fornecido pelo usu√°rio
2. Modo Prenota√ß√£o INCRA: Busca autom√°tica em rede e convers√£o TIFF‚ÜíPDF

Requisitos:
- pip install google-generativeai openpyxl python-docx pillow pdf2image --break-system-packages
"""

import os
import sys
import json
import shutil
import math
from pathlib import Path
from typing import Optional, Dict, List

# Importa√ß√µes das bibliotecas necess√°rias
try:
    import google.generativeai as genai
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image
    from pdf2image import convert_from_path
except ImportError as e:
    print(f"‚ùå Erro: Biblioteca necess√°ria n√£o encontrada - {e}")
    print("\nüì¶ Instale as depend√™ncias com:")
    print("pip install google-generativeai openpyxl python-docx pillow pdf2image --break-system-packages")
    sys.exit(1)


# ============================================================================
# CONFIGURA√á√ïES GLOBAIS
# ============================================================================

# Configura√ß√µes do INCRA
INCRA_CONFIG = {
    'base_path': r'\\192.168.20.100\trabalho\TRABALHO\IMAGENS\IMOVEIS\DOCUMENTOS - DIVERSOS',
    'folder_interval': 1000,  # Intervalo de agrupamento (1000 em 1000)
    'identificador_inicio': [
        'MINIST√âRIO DA AGRICULTURA, PECU√ÅRIA E ABASTECIMENTO',
        'INSTITUTO NACIONAL DE COLONIZA√á√ÉO E REFORMA AGR√ÅRIA',
        'MEMORIAL DESCRITIVO'
    ],
    'marcador_tabela': 'DESCRI√á√ÉO DA PARCELA',
    'marcador_azimutes': 'Azimutes: Azimutes geod√©sicos',
    'colunas_vertice': ['C√≥digo', 'Longitude', 'Latitude', 'Altitude'],
    'colunas_segmento': ['C√≥digo', 'Azimute', 'Dist.', 'Confronta√ß√µes']
}


# ============================================================================
# FUN√á√ïES AUXILIARES
# ============================================================================

def formatar_prenotacao(numero: str) -> str:
    """
    Formata n√∫mero de prenota√ß√£o para 8 d√≠gitos com zeros √† esquerda
    
    Args:
        numero: N√∫mero da prenota√ß√£o (com ou sem zeros √† esquerda)
    
    Returns:
        N√∫mero formatado com 8 d√≠gitos
    """
    # Remove espa√ßos e zeros √† esquerda, depois formata
    numero_limpo = str(int(numero.strip()))
    return numero_limpo.zfill(8)


def calcular_pasta_milhar(prenotacao: str) -> str:
    """
    Calcula a pasta de milhar onde o arquivo est√° armazenado
    
    Args:
        prenotacao: N√∫mero da prenota√ß√£o formatado (8 d√≠gitos)
    
    Returns:
        Nome da pasta de milhar (ex: '00230000')
    """
    numero = int(prenotacao)
    milhar_superior = math.ceil(numero / 1000) * 1000
    return str(milhar_superior).zfill(8)


def testar_acesso_rede() -> bool:
    """
    Testa se o caminho de rede do INCRA est√° acess√≠vel
    
    Returns:
        True se acess√≠vel, False caso contr√°rio
    """
    base_path = INCRA_CONFIG['base_path']
    
    print(f"\nüîå Testando acesso √† rede...")
    print(f"üìÇ Caminho: {base_path}")
    
    try:
        # Tenta acessar diretamente com os.scandir (mais compat√≠vel com UNC)
        with os.scandir(base_path) as entries:
            # Conta quantas pastas existem
            dirs = [entry.name for entry in entries if entry.is_dir()]
            
            print(f"‚úÖ Rede acess√≠vel!")
            print(f"üìÅ Encontradas {len(dirs)} pastas na rede")
            
            # Mostra algumas pastas como exemplo
            if dirs:
                exemplos = sorted(dirs)[:5]
                print(f"   Exemplos: {', '.join(exemplos)}")
            
            return True
            
    except PermissionError:
        print(f"‚ùå Acesso negado!")
        print(f"\nüí° Poss√≠veis causas:")
        print(f"   1. Sem permiss√µes de leitura")
        print(f"   2. Credenciais de rede necess√°rias")
        print(f"   3. Compartilhamento requer autentica√ß√£o")
        print(f"\nüîß Solu√ß√£o:")
        print(f"   Abra o Explorer e acesse primeiro:")
        print(f"   {base_path}")
        print(f"   Depois tente novamente o script.")
        return False
        
    except FileNotFoundError:
        print(f"‚ùå Caminho n√£o encontrado!")
        print(f"\nüí° Poss√≠veis causas:")
        print(f"   1. Servidor offline")
        print(f"   2. Caminho incorreto")
        print(f"   3. Rede desconectada")
        print(f"\nüîß Teste no CMD:")
        print(f"   dir \"{base_path}\"")
        return False
        
    except OSError as e:
        print(f"‚ùå Erro ao acessar rede: {e}")
        print(f"\nüí° Poss√≠veis causas:")
        print(f"   1. Timeout de rede")
        print(f"   2. Firewall bloqueando")
        print(f"   3. Protocolo SMB desabilitado")
        return False
    
    except Exception as e:
        print(f"‚ùå Erro inesperado: {e}")
        return False


def buscar_arquivo_incra(prenotacao: str) -> Optional[Path]:
    """
    Busca arquivo TIFF da prenota√ß√£o na rede do INCRA
    
    Args:
        prenotacao: N√∫mero da prenota√ß√£o (com ou sem formata√ß√£o)
    
    Returns:
        Path do arquivo se encontrado, None caso contr√°rio
    """
    # Formata prenota√ß√£o
    prenotacao_formatada = formatar_prenotacao(prenotacao)
    print(f"üîç Buscando prenota√ß√£o: {prenotacao_formatada}")
    
    # Calcula pasta de milhar
    pasta_milhar = calcular_pasta_milhar(prenotacao_formatada)
    print(f"üìÅ Pasta calculada: {pasta_milhar}")
    
    # Monta caminho completo como string (melhor para UNC)
    base_path = INCRA_CONFIG['base_path']
    
    # Garante que n√£o tem barra no final
    if base_path.endswith('\\') or base_path.endswith('/'):
        base_path = base_path[:-1]
    
    # Monta caminho da pasta e do arquivo
    pasta_completa = f"{base_path}\\{pasta_milhar}"
    arquivo_completo = f"{pasta_completa}\\{prenotacao_formatada}.tif"
    
    print(f"üìÇ Caminho: {arquivo_completo}")
    
    # Verifica se arquivo existe
    try:
        # M√©todo 1: Tenta acessar diretamente o arquivo
        if os.path.isfile(arquivo_completo):
            print(f"‚úÖ Arquivo encontrado!")
            return Path(arquivo_completo)
        
        # M√©todo 2: Se n√£o encontrou, lista a pasta para debug
        print(f"‚ùå Arquivo n√£o encontrado diretamente")
        
        if not os.path.isdir(pasta_completa):
            print(f"‚ùå Pasta n√£o existe: {pasta_milhar}")
            return None
        
        print(f"üìÅ Pasta existe, listando arquivos...")
        
        # Lista arquivos .tif na pasta
        arquivos_tif = []
        with os.scandir(pasta_completa) as entries:
            for entry in entries:
                if entry.is_file() and entry.name.lower().endswith('.tif'):
                    arquivos_tif.append(entry.name)
        
        if arquivos_tif:
            print(f"   Encontrados {len(arquivos_tif)} arquivos .tif na pasta")
            # Mostra alguns exemplos
            for arq in sorted(arquivos_tif)[:5]:
                print(f"   - {arq}")
            if len(arquivos_tif) > 5:
                print(f"   ... e mais {len(arquivos_tif) - 5} arquivos")
            
            # Procura o arquivo espec√≠fico na lista
            nome_procurado = f"{prenotacao_formatada}.tif"
            if nome_procurado.upper() in [a.upper() for a in arquivos_tif]:
                print(f"‚úÖ Arquivo encontrado na listagem!")
                return Path(arquivo_completo)
        else:
            print(f"   ‚ö†Ô∏è Pasta vazia ou sem arquivos .tif")
        
        return None
        
    except PermissionError as e:
        print(f"‚ùå Acesso negado: {e}")
        print(f"üí° Abra a pasta no Explorer primeiro:")
        print(f"   {pasta_completa}")
        return None
        
    except Exception as e:
        print(f"‚ùå Erro ao buscar arquivo: {e}")
        return None


def copiar_para_downloads(arquivo_origem: Path, prenotacao: str) -> Path:
    """
    Copia arquivo TIFF para pasta Tabelas_Incra em Documentos

    Args:
        arquivo_origem: Path do arquivo original (pode ser string UNC)
        prenotacao: N√∫mero da prenota√ß√£o formatado

    Returns:
        Path do arquivo copiado
    """
    # Determina pasta Documentos/Tabelas_Incra do usu√°rio
    home = Path.home()
    tabelas_incra = home / 'Documents' / 'Tabelas_Incra'

    # Cria pasta base se n√£o existir
    tabelas_incra.mkdir(parents=True, exist_ok=True)

    # Cria subpasta espec√≠fica para esta prenota√ß√£o
    pasta_prenotacao = tabelas_incra / f'Prenotacao_{prenotacao}'
    pasta_prenotacao.mkdir(parents=True, exist_ok=True)

    print(f"üìÅ Pasta criada: {pasta_prenotacao}")
    
    # Nome do arquivo
    nome_arquivo = os.path.basename(str(arquivo_origem))
    destino = pasta_prenotacao / nome_arquivo
    
    print(f"üìã Copiando arquivo...")
    print(f"   Origem: {arquivo_origem}")
    print(f"   Destino: {destino}")
    
    try:
        # Usa shutil.copy2 com strings para melhor compatibilidade UNC
        shutil.copy2(str(arquivo_origem), str(destino))
        print(f"‚úÖ Arquivo copiado com sucesso!")
    except Exception as e:
        print(f"‚ùå Erro ao copiar arquivo: {e}")
        print(f"‚ö†Ô∏è Tentando m√©todo alternativo...")
        
        # M√©todo alternativo: l√™ e escreve byte a byte
        try:
            with open(str(arquivo_origem), 'rb') as f_origem:
                conteudo = f_origem.read()
            with open(str(destino), 'wb') as f_destino:
                f_destino.write(conteudo)
            print(f"‚úÖ Arquivo copiado com m√©todo alternativo!")
        except Exception as e2:
            print(f"‚ùå Erro no m√©todo alternativo: {e2}")
            raise Exception(f"N√£o foi poss√≠vel copiar o arquivo: {e2}")
    
    return destino


def converter_tiff_para_pdf(tiff_path: Path) -> Path:
    """
    Converte arquivo TIFF (multi-p√°gina) para PDF
    
    Args:
        tiff_path: Path do arquivo TIFF
    
    Returns:
        Path do arquivo PDF gerado
    """
    print(f"\nüîÑ Convertendo TIFF para PDF...")
    
    pdf_path = tiff_path.with_suffix('.pdf')
    
    try:
        # Abre arquivo TIFF
        img = Image.open(tiff_path)
        
        # Lista para armazenar todas as p√°ginas
        images = []
        
        # Itera por todas as p√°ginas do TIFF
        try:
            page = 0
            while True:
                img.seek(page)
                # Converte para RGB se necess√°rio
                if img.mode != 'RGB':
                    rgb_img = img.convert('RGB')
                else:
                    rgb_img = img.copy()
                images.append(rgb_img)
                page += 1
                print(f"  üìÑ P√°gina {page} processada")
        except EOFError:
            pass  # Fim do arquivo TIFF
        
        # Salva como PDF multi-p√°gina
        if images:
            images[0].save(
                pdf_path,
                save_all=True,
                append_images=images[1:] if len(images) > 1 else [],
                resolution=100.0,
                quality=95
            )
            print(f"‚úÖ PDF criado: {pdf_path.name} ({len(images)} p√°ginas)")
            return pdf_path
        else:
            raise ValueError("Nenhuma p√°gina encontrada no TIFF")
            
    except Exception as e:
        print(f"‚ùå Erro ao converter TIFF: {e}")
        raise


def extrair_memorial_incra(pdf_path: Path, api_key: str) -> Dict:
    """
    Extrai tabela de Memorial Descritivo do INCRA
    
    Esta fun√ß√£o procura especificamente pelo formato do INCRA e extrai
    a tabela de coordenadas que pode estar em m√∫ltiplas p√°ginas
    
    Args:
        pdf_path: Path do arquivo PDF
        api_key: Chave da API do Gemini
    
    Returns:
        Dados estruturados da tabela
    """
    print(f"\nüìä Extraindo Memorial Descritivo do INCRA...")
    
    # Configura API
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash-lite')
    
    # Carrega PDF
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()
    
    # Prompt especializado para Memorial do INCRA
    prompt = f"""
Voc√™ est√° processando um Memorial Descritivo do INCRA (Instituto Nacional de Coloniza√ß√£o e Reforma Agr√°ria).

INSTRU√á√ïES CR√çTICAS:

1. LOCALIZA√á√ÉO: Encontre o bloco que cont√©m:
   - "MINIST√âRIO DA AGRICULTURA, PECU√ÅRIA E ABASTECIMENTO"
   - "INSTITUTO NACIONAL DE COLONIZA√á√ÉO E REFORMA AGR√ÅRIA"
   - "MEMORIAL DESCRITIVO"

2. TABELA: Ap√≥s encontrar "DESCRI√á√ÉO DA PARCELA" e "Azimutes: Azimutes geod√©sicos", 
   localize a tabela com as seguintes colunas:

   V√âRTICE (4 colunas):
   - C√≥digo
   - Longitude
   - Latitude  
   - Altitude (m)

   SEGMENTO VANTE (4 colunas):
   - C√≥digo
   - Azimute
   - Dist. (m)
   - Confronta√ß√µes

3. MULTI-P√ÅGINA: A tabela pode continuar em m√∫ltiplas p√°ginas. Continue extraindo 
   at√© encontrar um novo cabe√ßalho de se√ß√£o (como "CERTIFICA√á√ÉO") ou o fim da tabela.

4. FORMATO DE SA√çDA: Retorne APENAS o JSON neste formato exato:

{{
  "header_row1": ["V√âRTICE", "SEGMENTO VANTE"],
  "header_row2": ["C√≥digo", "Longitude", "Latitude", "Altitude (m)", "C√≥digo", "Azimute", "Dist. (m)", "Confronta√ß√µes"],
  "data": [
    ["valor1", "valor2", "valor3", "valor4", "valor5", "valor6", "valor7", "valor8"],
    ...
  ]
}}

IMPORTANTE:
- Mantenha a formata√ß√£o exata dos valores (graus, aspas, v√≠rgulas)
- Se um campo estiver vazio, use ""
- Extraia TODAS as linhas da tabela de TODAS as p√°ginas
- Retorne APENAS o JSON, sem texto adicional
"""
    
    print("ü§ñ Enviando para Gemini API...")
    response = model.generate_content([
        prompt,
        {"mime_type": "application/pdf", "data": pdf_data}
    ])
    
    print("‚úÖ Resposta recebida")
    
    # Processa resposta
    response_text = response.text.strip()
    
    # Remove marcadores markdown se presentes
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
    
    # Extrai JSON
    if '{' in response_text:
        response_text = response_text[response_text.find('{'):]
    if '}' in response_text:
        response_text = response_text[:response_text.rfind('}')+1]
    
    response_text = response_text.strip()
    
    # Parse JSON
    try:
        table_data = json.loads(response_text)
        num_linhas = len(table_data.get('data', []))
        print(f"‚úÖ Tabela extra√≠da: {num_linhas} linhas de dados")
        return table_data
    except json.JSONDecodeError as e:
        print(f"‚ùå Erro ao fazer parse do JSON: {e}")
        print(f"Resposta recebida (primeiros 500 chars): {response_text[:500]}")
        raise


# ============================================================================
# FUN√á√ïES PRINCIPAIS DO MODO NORMAL
# ============================================================================

def configure_gemini_api(api_key):
    """Configura a API do Google Gemini"""
    genai.configure(api_key=api_key)


def extract_table_from_pdf(pdf_path, api_key):
    """Extrai dados da tabela do PDF usando Google Gemini API (modo normal)"""
    print(f"üìÑ Processando PDF: {pdf_path}")
    
    configure_gemini_api(api_key)
    model = genai.GenerativeModel('gemini-2.5-flash-lite')
    
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()
    
    prompt = """
Analise este Memorial Descritivo e extraia APENAS a tabela principal que cont√©m informa√ß√µes de v√©rtices.

A tabela tem a seguinte estrutura:
- Cabe√ßalho Linha 1: "V√âRTICE" (colunas A-D) e "SEGMENTO VANTE" (colunas E-H)
- Cabe√ßalho Linha 2: C√≥digo, Longitude, Latitude, Altitude (m), C√≥digo, Azimute, Dist. (m), Confronta√ß√µes

Retorne os dados em formato JSON seguindo EXATAMENTE esta estrutura:
{
  "header_row1": ["V√âRTICE", "SEGMENTO VANTE"],
  "header_row2": ["C√≥digo", "Longitude", "Latitude", "Altitude (m)", "C√≥digo", "Azimute", "Dist. (m)", "Confronta√ß√µes"],
  "data": [
    ["valor1", "valor2", "valor3", "valor4", "valor5", "valor6", "valor7", "valor8"],
    ...
  ]
}

IMPORTANTE: 
- Retorne APENAS o JSON, sem texto adicional
- Inclua TODOS os dados da tabela
- Mantenha a formata√ß√£o original dos valores
- Se um campo estiver vazio, use ""
"""
    
    print("ü§ñ Enviando para Gemini API...")
    response = model.generate_content([prompt, {"mime_type": "application/pdf", "data": pdf_data}])
    
    print("‚úÖ Resposta recebida da API")
    
    response_text = response.text.strip()
    
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
    
    response_text = response_text.strip()
    
    try:
        table_data = json.loads(response_text)
        print(f"üìä Tabela extra√≠da: {len(table_data.get('data', []))} linhas de dados")
        return table_data
    except json.JSONDecodeError as e:
        print(f"‚ùå Erro ao fazer parse do JSON: {e}")
        print(f"Resposta recebida: {response_text[:500]}...")
        sys.exit(1)


def create_excel_file(table_data, output_path):
    """Cria arquivo Excel com a tabela formatada"""
    print(f"\nüìä Criando arquivo Excel: {output_path}")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Memorial Descritivo"
    
    header_font = Font(bold=True, size=11)
    center_alignment = Alignment(horizontal='center', vertical='center')
    border_style = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Linha 1: Cabe√ßalhos mesclados
    ws.merge_cells('A1:D1')
    cell_a1 = ws['A1']
    cell_a1.value = "V√âRTICE"
    cell_a1.font = header_font
    cell_a1.alignment = center_alignment
    cell_a1.border = border_style
    
    ws.merge_cells('E1:H1')
    cell_e1 = ws['E1']
    cell_e1.value = "SEGMENTO VANTE"
    cell_e1.font = header_font
    cell_e1.alignment = center_alignment
    cell_e1.border = border_style
    
    # Linha 2: Sub-cabe√ßalhos
    header_row2 = table_data.get('header_row2', [])
    for col_idx, header in enumerate(header_row2, start=1):
        cell = ws.cell(row=2, column=col_idx)
        cell.value = header
        cell.font = header_font
        cell.alignment = center_alignment
        cell.border = border_style
    
    # Linhas 3+: Dados
    data_rows = table_data.get('data', [])
    for row_idx, row_data in enumerate(data_rows, start=3):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = value
            cell.border = border_style
            if col_idx in [1, 5]:
                cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Ajusta larguras
    column_widths = {
        'A': 15, 'B': 18, 'C': 18, 'D': 15,
        'E': 15, 'F': 15, 'G': 12, 'H': 30
    }
    
    for col_letter, width in column_widths.items():
        ws.column_dimensions[col_letter].width = width
    
    wb.save(output_path)
    print(f"‚úÖ Excel criado com sucesso!")
    return output_path


def create_word_file(table_data, output_path):
    """Cria arquivo Word com a tabela formatada"""
    print(f"\nüìù Criando arquivo Word: {output_path}")
    
    doc = Document()

    data_rows = table_data.get('data', [])
    num_rows = 2 + len(data_rows)
    num_cols = 8
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Linha 1: Mesclagem
    cell_a1 = table.rows[0].cells[0]
    cell_d1 = table.rows[0].cells[3]
    cell_a1.merge(cell_d1)
    cell_a1.text = "V√âRTICE"
    cell_a1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_a1.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(11)
    
    cell_e1 = table.rows[0].cells[4]
    cell_h1 = table.rows[0].cells[7]
    cell_e1.merge(cell_h1)
    cell_e1.text = "SEGMENTO VANTE"
    cell_e1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_e1.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(11)
    
    # Linha 2: Sub-cabe√ßalhos
    header_row2 = table_data.get('header_row2', [])
    for col_idx, header in enumerate(header_row2):
        cell = table.rows[1].cells[col_idx]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
    
    # Linhas 3+: Dados
    for row_idx, row_data in enumerate(data_rows, start=2):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(value) if value else ""
            if col_idx in [0, 4]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    # AutoFit
    table.autofit = True
    table.allow_autofit = True
    
    # Larguras preferenciais
    preferred_widths = [
        Cm(2.5), Cm(3.5), Cm(3.5), Cm(2.5),
        Cm(2.5), Cm(2.5), Cm(2.0), Cm(4.0)
    ]
    
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(preferred_widths):
                cell.width = preferred_widths[idx]
    
    doc.save(output_path)
    print(f"‚úÖ Word criado com sucesso!")
    return output_path


# ============================================================================
# FUN√á√ÉO PRINCIPAL - MODO PRENOTA√á√ÉO INCRA
# ============================================================================

def modo_prenotacao_incra(api_key: str):
    """
    Modo de opera√ß√£o para prenota√ß√µes do INCRA
    
    Fluxo completo:
    1. Testa acesso √† rede
    2. Solicita n√∫mero da prenota√ß√£o
    3. Busca arquivo TIFF na rede
    4. Copia para Downloads
    5. Converte TIFF para PDF
    6. Extrai tabela do Memorial
    7. Oferece op√ß√£o de gerar Excel e/ou Word
    """
    print("\n" + "="*70)
    print("üèõÔ∏è  MODO PRENOTA√á√ÉO INCRA")
    print("="*70)
    
    # 0. Testa acesso √† rede primeiro
    if not testar_acesso_rede():
        print("\n" + "="*70)
        print("‚ùå ERRO: N√£o foi poss√≠vel acessar a rede do INCRA!")
        print("="*70)
        print("\nüìù Verifique:")
        print("  1. Conex√£o com a rede")
        print("  2. Caminho configurado (linha ~30 do script)")
        print(f"     Atual: {INCRA_CONFIG['base_path']}")
        print("  3. Permiss√µes de acesso")
        print("  4. VPN ativa (se necess√°rio)")
        return
    
    # 1. Solicita prenota√ß√£o
    prenotacao = input("\nüìã Digite o n√∫mero da Prenota√ß√£o (ex: 229885 ou 00229885): ").strip()
    
    if not prenotacao:
        print("‚ùå N√∫mero de prenota√ß√£o n√£o fornecido!")
        return
    
    try:
        prenotacao_formatada = formatar_prenotacao(prenotacao)
    except ValueError:
        print("‚ùå N√∫mero de prenota√ß√£o inv√°lido!")
        return
    
    print(f"‚úÖ Prenota√ß√£o formatada: {prenotacao_formatada}")
    
    # 2. Busca arquivo TIFF
    print("\n" + "-"*70)
    arquivo_tiff = buscar_arquivo_incra(prenotacao_formatada)
    
    if not arquivo_tiff:
        print("‚ùå Arquivo n√£o encontrado na rede do INCRA!")
        print(f"üìÇ Caminho esperado: {INCRA_CONFIG['base_path']}\\{calcular_pasta_milhar(prenotacao_formatada)}\\{prenotacao_formatada}.tif")
        return
    
    # 3. Copia para Downloads
    print("\n" + "-"*70)
    arquivo_local = copiar_para_downloads(arquivo_tiff, prenotacao_formatada)
    
    # 4. Converte TIFF para PDF
    print("\n" + "-"*70)
    try:
        arquivo_pdf = converter_tiff_para_pdf(arquivo_local)
    except Exception as e:
        print(f"‚ùå Erro na convers√£o: {e}")
        return
    
    # 5. Extrai tabela
    print("\n" + "-"*70)
    try:
        table_data = extrair_memorial_incra(arquivo_pdf, api_key)
    except Exception as e:
        print(f"‚ùå Erro na extra√ß√£o: {e}")
        return
    
    # 6. Oferece op√ß√µes de gera√ß√£o
    print("\n" + "="*70)
    print("üìä EXTRA√á√ÉO CONCLU√çDA!")
    print(f"‚úÖ {len(table_data.get('data', []))} linhas extra√≠das")
    print("="*70)
    
    escolher_arquivos_saida(table_data, arquivo_pdf.parent, prenotacao_formatada)


# ============================================================================
# FUN√á√ÉO PRINCIPAL - MODO NORMAL
# ============================================================================

def modo_normal(api_key: str):
    """Modo de opera√ß√£o normal (arquivo PDF fornecido pelo usu√°rio)"""
    print("\n" + "="*70)
    print("üìÑ MODO NORMAL - Processar PDF")
    print("="*70)
    
    pdf_path = input("\nüìÇ Digite o caminho completo do arquivo PDF: ").strip()
    pdf_path = pdf_path.strip("'\"")
    
    if not os.path.exists(pdf_path):
        print(f"\n‚ùå Erro: Arquivo n√£o encontrado: {pdf_path}")
        return
    
    try:
        table_data = extract_table_from_pdf(pdf_path, api_key)
    except Exception as e:
        print(f"\n‚ùå Erro ao processar PDF: {e}")
        return
    
    output_dir = Path(pdf_path).parent
    
    print("\n" + "="*70)
    print("üìä EXTRA√á√ÉO CONCLU√çDA!")
    print(f"‚úÖ {len(table_data.get('data', []))} linhas extra√≠das")
    print("="*70)
    
    escolher_arquivos_saida(table_data, output_dir)


# ============================================================================
# FUN√á√ÉO DE ESCOLHA DE ARQUIVOS DE SA√çDA
# ============================================================================

def escolher_arquivos_saida(table_data: Dict, output_dir: Path, prefixo: str = "output"):
    """
    Permite ao usu√°rio escolher quais arquivos gerar (Excel e/ou Word)
    
    Args:
        table_data: Dados extra√≠dos da tabela
        output_dir: Diret√≥rio onde salvar os arquivos
        prefixo: Prefixo para nome dos arquivos
    """
    print("\n" + "="*70)
    print("üíæ ESCOLHA OS ARQUIVOS DE SA√çDA")
    print("="*70)
    print("\nQuais arquivos voc√™ deseja gerar?")
    print("  1 - Apenas Excel (.xlsx)")
    print("  2 - Apenas Word (.docx)")
    print("  3 - Ambos (Excel + Word)")
    print("  0 - Cancelar (n√£o gerar nenhum)")
    
    while True:
        escolha = input("\nüëâ Digite sua escolha (0-3): ").strip()
        
        if escolha == '0':
            print("\n‚ùå Opera√ß√£o cancelada. Nenhum arquivo foi gerado.")
            return
        
        elif escolha == '1':
            # Apenas Excel
            excel_path = output_dir / f"{prefixo}.xlsx"
            create_excel_file(table_data, str(excel_path))
            print(f"\n‚úÖ Arquivo gerado:")
            print(f"   üìä {excel_path}")
            break
        
        elif escolha == '2':
            # Apenas Word
            word_path = output_dir / f"{prefixo}.docx"
            create_word_file(table_data, str(word_path))
            print(f"\n‚úÖ Arquivo gerado:")
            print(f"   üìù {word_path}")
            break
        
        elif escolha == '3':
            # Ambos
            excel_path = output_dir / f"{prefixo}.xlsx"
            word_path = output_dir / f"{prefixo}.docx"
            create_excel_file(table_data, str(excel_path))
            create_word_file(table_data, str(word_path))
            print(f"\n‚úÖ Arquivos gerados:")
            print(f"   üìä {excel_path}")
            print(f"   üìù {word_path}")
            break
        
        else:
            print("‚ùå Op√ß√£o inv√°lida! Digite 0, 1, 2 ou 3.")


# ============================================================================
# FUN√á√ÉO MAIN
# ============================================================================

def main():
    """Fun√ß√£o principal que coordena o fluxo completo"""
    print("="*70)
    print("üöÄ PROCESSADOR DE MEMORIAL DESCRITIVO")
    print("="*70)
    
    # Configurar API Key (fixa)
    print("\nüîë Configura√ß√£o da API do Google Gemini")
    api_key = 'AIzaSyAdA_GO7cQ0m1ouie4wGwXf4a4SnHKjBh8'
    print(f"‚úÖ Usando chave configurada")
    
    # Escolher modo de opera√ß√£o
    print("\n" + "="*70)
    print("üéØ ESCOLHA O MODO DE OPERA√á√ÉO")
    print("="*70)
    print("\n  1 - Modo Normal (fornecer arquivo PDF)")
    print("  2 - Modo Prenota√ß√£o INCRA (busca autom√°tica)")
    
    while True:
        modo = input("\nüëâ Digite sua escolha (1 ou 2): ").strip()
        
        if modo == '1':
            modo_normal(api_key)
            break
        elif modo == '2':
            modo_prenotacao_incra(api_key)
            break
        else:
            print("‚ùå Op√ß√£o inv√°lida! Digite 1 ou 2.")
    
    print("\n" + "="*70)
    print("‚ú® PROCESSAMENTO FINALIZADO!")
    print("="*70)


if __name__ == "__main__":
    main()